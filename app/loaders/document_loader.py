"""
Step 1: Document Loading
-------------------------
Goal: given a file path, return clean plain text + minimal metadata.

Why this matters for RAG:
Every downstream component (chunking, embedding, retrieval) inherits
whatever quality of text comes out of here. A PDF that extracts as
garbled text with broken word-spacing will produce bad embeddings,
which produce bad retrieval, which produces bad LLM answers — and
none of those later stages will throw an error. The bug is silent.

Design decision: we return a simple `Document` object (path, text,
metadata) rather than raw strings, because we'll want to carry
metadata (source filename, page number) all the way through to the
final answer, so we can cite sources later.
"""

from dataclasses import dataclass, field
from pathlib import Path
from typing import Any


@dataclass
class Document:
    text: str
    source: str                      # filename, for citation later
    metadata: dict[str, Any] = field(default_factory=dict)


def load_txt(path: Path) -> Document:
    text = path.read_text(encoding="utf-8", errors="ignore")
    return Document(text=text, source=path.name)


def load_pdf(path: Path) -> Document:
    # pdfplumber over PyPDF2: better at preserving layout/word-spacing,
    # which matters a lot for extraction quality on real-world PDFs.
    import pdfplumber

    pages_text = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            page_text = page.extract_text() or ""
            pages_text.append(page_text)

    full_text = "\n\n".join(pages_text)
    return Document(
        text=full_text,
        source=path.name,
        metadata={"num_pages": len(pages_text)},
    )


def load_docx(path: Path) -> Document:
    import docx  # python-docx

    doc = docx.Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = "\n\n".join(paragraphs)
    return Document(text=full_text, source=path.name)


LOADERS = {
    ".txt": load_txt,
    ".md": load_txt,
    ".pdf": load_pdf,
    ".docx": load_docx,
}


def load_document(path: str | Path) -> Document:
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix not in LOADERS:
        raise ValueError(
            f"Unsupported file type: {suffix}. "
            f"Supported: {list(LOADERS.keys())}"
        )
    return LOADERS[suffix](path)


if __name__ == "__main__":
    # Quick manual test
    import sys
    doc = load_document(sys.argv[1])
    print(f"Loaded {doc.source}: {len(doc.text)} chars")
    print(doc.metadata)
    print("---first 300 chars---")
    print(doc.text[:300])
